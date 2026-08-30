"""Document reader module.

Reads text from Microsoft Word documents (DOC, DOCX).
Uses python-docx for DOCX and antiword for legacy DOC files.
"""

import subprocess
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, Any

from config import settings
from utils.logger import setup_logger


logger = setup_logger(__name__)


def read_docx(file_path: Path, file_id: str) -> Dict[str, Any]:
    """Read DOCX file using python-docx.
    
    Extracts text from paragraphs and tables.
    
    Args:
        file_path: Path to DOCX file
        file_id: File identifier
    
    Returns:
        Dictionary with extracted text and metadata
    """
    import docx
    
    doc = docx.Document(file_path)
    paragraphs = []
    
    # Extract paragraph text
    for p in doc.paragraphs:
        if p.text:
            paragraphs.append(p.text)
    
    # Extract table text
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                paragraphs.append(' | '.join(row_text))
    
    full_text = '\n'.join(paragraphs)
    page_count = max(1, len(full_text) // 3000)  # Estimate pages
    
    return {
        'id': file_id,
        'status': 'done',
        'type_read': 'docx',
        'page_count': page_count,
        'char_count': len(full_text),
        'image_quality': None,
        'full_text': full_text
    }


def read_doc_antiword(file_path: Path, file_id: str) -> Dict[str, Any]:
    """Read legacy DOC file using antiword.
    
    Args:
        file_path: Path to DOC file
        file_id: File identifier
    
    Returns:
        Dictionary with extracted text and metadata
    """
    antiword_path = settings.antiword.ANTIWORD_PATH
    antiword_home = settings.antiword.ANTIWORD_HOME
    
    env = {}
    if antiword_home:
        env['ANTIWORDHOME'] = antiword_home
    
    result = subprocess.run(
        [antiword_path, str(file_path)],
        capture_output=True,
        text=True,
        timeout=30,
        cwd=antiword_home if antiword_home else None,
        env=env or None
    )
    result.check_returncode()
    
    full_text = result.stdout
    page_count = max(1, len(full_text) // 3000)
    
    return {
        'id': file_id,
        'status': 'done',
        'type_read': 'doc',
        'page_count': page_count,
        'char_count': len(full_text),
        'image_quality': None,
        'full_text': full_text
    }


def read_docx_zip(file_path: Path, file_id: str) -> Dict[str, Any]:
    """Read DOCX/DOC directly from ZIP archive as fallback.
    
    Bypasses python-docx and reads XML directly.
    
    Args:
        file_path: Path to file
        file_id: File identifier
    
    Returns:
        Dictionary with extracted text and metadata
    """
    with zipfile.ZipFile(file_path, 'r') as z:
        if 'word/document.xml' in z.namelist():
            xml_content = z.read('word/document.xml')
        elif 'word/document2.xml' in z.namelist():
            xml_content = z.read('word/document2.xml')
        else:
            raise ValueError("word/document.xml not found in archive")
    
    root = ET.fromstring(xml_content)
    
    # Remove namespaces from tags
    for elem in root.iter():
        if '}' in elem.tag:
            elem.tag = elem.tag.split('}', 1)[1]
    
    # Extract text from paragraphs and tables
    paragraphs = []
    for p in root.iter('p'):
        texts = []
        for t in p.iter('t'):
            if t.text:
                texts.append(t.text)
        if texts:
            paragraphs.append(''.join(texts))
    
    full_text = '\n'.join(paragraphs)
    page_count = max(1, len(full_text) // 3000)
    
    return {
        'id': file_id,
        'status': 'done',
        'type_read': 'docx',
        'page_count': page_count,
        'char_count': len(full_text),
        'image_quality': None,
        'full_text': full_text
    }
